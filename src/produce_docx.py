import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

txt_dir = '../txt/'

CHORDS = ['DO', 'RE', 'MI', 'FA', 'SOL', 'LA', 'SI', 
          'DO7', 'RE7', 'MI7', 'FA7', 'SOL7', 'LA7', 'SI7', 
          'DOb', 'REb', 'MIb', 'FAb', 'SOLb', 'LAb', 'SIb',
          'DO#', 'RE#', 'MI#', 'FA#', 'SOL#', 'LA#', 'SI#',
          'DO#7', 'RE#7', 'MI#7', 'FA#7', 'SOL#7', 'LA#7', 'SI#7',]

def parse_song(content):
    song = {
        'title': '',
        'lyrics': []
    }
    lines = content.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            song['title'] = line
            continue
        if i == 1:
            continue
        chords_found = False
        for chord in CHORDS:
            if chord in line:
                chords_found = True
        song["lyrics"].append([chords_found, line])
    return song

songs = []

for filename in sorted(os.listdir(txt_dir)):
    if filename.endswith('.txt'):
        with open(os.path.join(txt_dir, filename), 'r') as txt_file:
            content = txt_file.read()
            song = parse_song(content)
            songs.append(song)

# Create a new document
doc = Document()
styles = doc.styles
chords_style = styles.add_style('Chords', WD_STYLE_TYPE.PARAGRAPH)
chords_style.hidden = False
chords_style.quick_style = True
# chords_style.base_style = styles['Normal']
chords_style.font.bold = True
chords_style.font.size = Pt(12)

lyrics_style = styles.add_style('Lyrics', WD_STYLE_TYPE.PARAGRAPH)
lyrics_style.hidden = False
lyrics_style.quick_style = True
# lyrics_style.base_style = styles['Normal']
lyrics_style.font.size = Pt(12)

# Add a title
title = doc.add_heading('CANZONIERE', level=1)

song_number = 0

for song in songs:
    song_number += 1
    title = str(song_number) + " - " + song['title']
    doc.add_heading(title, level=2)
    for is_chords, line in song["lyrics"]:
        paragraph = doc.add_paragraph(line)
        if is_chords:
            paragraph.style = chords_style
        else:
            paragraph.style = lyrics_style
    doc.add_page_break()
# Save the document
doc.save('canzoniere.docx')